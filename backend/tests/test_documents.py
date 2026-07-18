from __future__ import annotations

import base64
import io
import sqlite3
import threading
from pathlib import Path

import pytest
from docx import Document
from reportlab.pdfgen import canvas


@pytest.fixture
async def document_service(tmp_path: Path):
    from big_bear_ai.config import load_settings
    from big_bear_ai.database import Database
    from big_bear_ai.services.documents import DocumentService

    settings = load_settings(
        {
            "BIG_BEAR_DATA_DIR": str(tmp_path),
            "BIG_BEAR_MAX_UPLOAD_MB": "1",
            "LANGSMITH_TRACING": "false",
        }
    )
    database = Database(settings.database_path)
    await database.initialize()
    return DocumentService(database, settings)


def encoded(content: bytes) -> str:
    return base64.b64encode(content).decode("ascii")


@pytest.mark.asyncio
async def test_upload_indexes_searches_and_deletes_text(document_service) -> None:
    uploaded = await document_service.upload(
        {
            "filename": "contract.md",
            "title": "Contract Guide",
            "content_base64": encoded(
                b"API contract verification checks status codes and response schemas."
            ),
        }
    )

    assert uploaded["filename"] == "contract.md"
    assert uploaded["index_status"] == "ready"
    assert uploaded["chunk_count"] == 1
    stored_path = Path(uploaded["file_path"])
    assert stored_path.exists()

    results = await document_service.search("response schemas")
    assert results[0]["document_id"] == uploaded["id"]
    assert "response schemas" in results[0]["content"]

    await document_service.delete(uploaded["id"])
    assert not stored_path.exists()
    assert await document_service.get(uploaded["id"]) is None


@pytest.mark.asyncio
@pytest.mark.parametrize(
    ("filename", "content"),
    [
        ("notes.txt", b"plain text knowledge"),
        ("notes.md", b"# Markdown knowledge"),
        ("data.json", b'{"topic": "json knowledge"}'),
        ("data.csv", b"topic,value\ncsv,knowledge\n"),
    ],
)
async def test_upload_supports_utf8_text_formats(
    document_service, filename: str, content: bytes
) -> None:
    uploaded = await document_service.upload(
        {"filename": filename, "content_base64": encoded(content)}
    )
    assert uploaded["chunk_count"] == 1


@pytest.mark.asyncio
async def test_upload_extracts_real_pdf_and_docx(document_service) -> None:
    pdf_buffer = io.BytesIO()
    pdf = canvas.Canvas(pdf_buffer)
    pdf.drawString(72, 720, "API contract verification from PDF")
    pdf.save()

    docx_buffer = io.BytesIO()
    docx = Document()
    docx.add_paragraph("Mobile regression knowledge from DOCX")
    docx.save(docx_buffer)

    pdf_document = await document_service.upload(
        {"filename": "guide.pdf", "content_base64": encoded(pdf_buffer.getvalue())}
    )
    docx_document = await document_service.upload(
        {"filename": "guide.docx", "content_base64": encoded(docx_buffer.getvalue())}
    )

    assert pdf_document["chunk_count"] == 1
    assert docx_document["chunk_count"] == 1
    assert (await document_service.search("contract verification"))[0][
        "document_id"
    ] == pdf_document["id"]
    assert (await document_service.search("Mobile regression"))[0][
        "document_id"
    ] == docx_document["id"]


@pytest.mark.asyncio
async def test_search_matches_chinese_substrings(document_service) -> None:
    uploaded = await document_service.upload(
        {
            "filename": "chinese.txt",
            "content_base64": encoded("这是测试规范文档，包含边界分析。".encode("utf-8")),
        }
    )

    results = await document_service.search("边界")

    assert uploaded["id"] in {result["document_id"] for result in results}


@pytest.mark.asyncio
async def test_fresh_database_starts_without_documents(document_service) -> None:
    page = await document_service.list(limit=10)
    results = await document_service.search("测试流程")

    assert page == {"items": [], "total": 0, "next_cursor": None}
    assert results == []


@pytest.mark.asyncio
async def test_docx_extraction_includes_table_content(document_service) -> None:
    buffer = io.BytesIO()
    document = Document()
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Scenario"
    table.cell(0, 1).text = "matrix coverage"
    document.save(buffer)

    uploaded = await document_service.upload(
        {"filename": "matrix.docx", "content_base64": encoded(buffer.getvalue())}
    )

    assert (await document_service.search("matrix coverage"))[0][
        "document_id"
    ] == uploaded["id"]


@pytest.mark.asyncio
async def test_upload_creates_directories_off_the_event_loop(
    document_service, monkeypatch: pytest.MonkeyPatch
) -> None:
    original_mkdir = Path.mkdir
    event_loop_thread = threading.current_thread()

    def checked_mkdir(path: Path, *args, **kwargs):
        assert threading.current_thread() is not event_loop_thread
        return original_mkdir(path, *args, **kwargs)

    monkeypatch.setattr(Path, "mkdir", checked_mkdir)

    uploaded = await document_service.upload(
        {
            "filename": "non-blocking.txt",
            "content_base64": encoded(b"event loop safe upload"),
        }
    )

    assert uploaded["filename"] == "non-blocking.txt"


@pytest.mark.asyncio
@pytest.mark.parametrize(
    ("payload", "code"),
    [
        ({"filename": "../secret.txt", "content_base64": encoded(b"x")}, "INVALID_FILE"),
        ({"filename": "script.exe", "content_base64": encoded(b"x")}, "UNSUPPORTED_FILE"),
        ({"filename": "bad.txt", "content_base64": "%%%"}, "INVALID_FILE"),
        ({"filename": "huge.txt", "content_base64": encoded(b"x" * (1024 * 1024 + 1))}, "FILE_TOO_LARGE"),
    ],
)
async def test_upload_rejects_invalid_files(document_service, payload, code) -> None:
    from big_bear_ai.services.documents import DocumentError

    with pytest.raises(DocumentError) as error:
        await document_service.upload(payload)
    assert error.value.code == code


@pytest.mark.asyncio
async def test_failed_extraction_leaves_no_rows_or_files(document_service) -> None:
    from big_bear_ai.services.documents import DocumentError

    with pytest.raises(DocumentError) as error:
        await document_service.upload(
            {"filename": "broken.pdf", "content_base64": encoded(b"not a pdf")}
        )
    assert error.value.code == "EXTRACTION_FAILED"

    def count(connection: sqlite3.Connection) -> int:
        return connection.execute(
            "SELECT COUNT(*) FROM documents WHERE read_only = 0"
        ).fetchone()[0]

    assert await document_service.database.run(count) == 0
    assert not list(document_service.settings.uploads_dir.glob("*"))


@pytest.mark.asyncio
async def test_read_only_document_cannot_be_deleted(document_service) -> None:
    from big_bear_ai.repositories.resources import ReadOnlyResourceError

    def insert_read_only(connection: sqlite3.Connection) -> None:
        connection.execute(
            """
            INSERT INTO documents(
                id, title, description, filename, media_type, size_bytes,
                extracted_text, index_status, file_path, author, read_only,
                created_at, updated_at
            ) VALUES (
                'read-only-document', 'Read-only', '', 'readonly.txt', 'text/plain', 4,
                'text', 'ready', NULL, '系统', 1,
                '2026-07-17T00:00:00+00:00', '2026-07-17T00:00:00+00:00'
            )
            """
        )

    await document_service.database.run(insert_read_only)
    stored = await document_service.get("read-only-document")

    assert stored["read_only"] is True
    with pytest.raises(ReadOnlyResourceError):
        await document_service.delete("read-only-document")


@pytest.mark.asyncio
async def test_management_graph_exposes_document_workflow(tmp_path: Path) -> None:
    from big_bear_ai.config import load_settings
    from big_bear_ai.graphs.management import build_management_graph

    settings = load_settings(
        {"BIG_BEAR_DATA_DIR": str(tmp_path), "LANGSMITH_TRACING": "false"}
    )
    graph = build_management_graph(settings)

    uploaded = await graph.ainvoke(
        {
            "operation": "action",
            "resource": "document",
            "payload": {
                "action": "upload",
                "filename": "guide.txt",
                "content_base64": encoded(b"boundary value analysis guide"),
            },
        }
    )
    assert uploaded["ok"] is True

    searched = await graph.ainvoke(
        {
            "operation": "action",
            "resource": "document",
            "payload": {"action": "search", "query": "boundary value"},
        }
    )
    assert searched["data"][0]["document_id"] == uploaded["data"]["id"]

    deleted = await graph.ainvoke(
        {
            "operation": "delete",
            "resource": "document",
            "resource_id": uploaded["data"]["id"],
        }
    )
    assert deleted["ok"] is True
